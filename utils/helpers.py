from docx import Document
from PyPDF2 import PdfReader
import pandas as pd
import comtypes.client

def docx_to_text(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            row_text = ""
            for cell in row.cells:
                if cell.text:
                    row_text += cell.text + " "
            parts.append(row_text)
    return "\n".join(parts).strip()


def pdf_to_text(path: str) -> str:
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n".join(parts).strip()

def excel_to_text(path: str) -> str:
    table = pd.read_excel(path)
    return table.to_string(index=False)


def pptx_to_pdf(path: str) -> None:
    powerpoint = comtypes.client.CreateObject("PowerPoint.Application")
    powerpoint.Visible = 1
    presentation = powerpoint.Presentations.Open(path)
    presentation.SaveAs(path.replace(".pptx", ".pdf"), 32)
    presentation.Close()
    powerpoint.Quit()